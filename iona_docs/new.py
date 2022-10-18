# from docx import Document
# import docx
# import pyexcel
# import re
# import string


# def analyze(docfile):
#     doc = docx.Document(docfile)
#     tot = 0 
#     frequency = {}
#     freq = 0
#     for para in doc.paragraphs:
#         data = para.text

#         #splitting data into words

#         text = data.lower().split(" ")
#         #print(words)
#         match_pattern = re.findall(r's*([A-Z]|[a-z]|\s|\/|\t|[0-9]|\t|\.|\"|\,|\-|\_|\!|\;|\?)*\n?\s*', data)
#         for word in text: 
#             for word in match_pattern:
#                 count = frequency.get(word,0)
#                 frequency[word] = count + 1

#             frequency_list = frequency.keys()
#             for words in frequency_list:
#                 #print (words, frequency[words])
#             #print(word)  
#                 tot = tot + len(word)
                
#                 freq = frequency[words]/tot
#                 if(freq>0.001):
#                     print (words,":", freq)
#         #print(tot) 


# analyze("C:/Users/ionad/OneDrive - University College Cork/Desktop/UCC/SEM 2/CS 6507 - Prog in Python with DS Apps/Assignment/2/pride_and_prejudice.docx")






import docx
import pyexcel 
import re
import xlsxwriter



def analyze(docfile):
    doc = docx.Document(docfile)
    data = []
    d= dict()
    freq = 0
    count =0
    tot = 0
    result ={}
    punctuation = ".?:,';()\.-_"

    for para in doc.paragraphs:
        data.append(para.text)
    #pat = (re.findall(r'(s*([A-Z]|[a-z]|\s|\/|\t|[0-9]|\t|\.|\"|\,|\-|\_|\!|\;|\?)*\n)\s*', '\n'.join(newparatextlist))) 
    #pat = (re.findall(r'([a-z]|[A-Z]|[0-9]|\/|\')*', '\n'.join(newparatextlist)))  
    #print(newparatextlist)
    for line in data:
        words = line.strip(" ").split(" ")
        count =count +len(words) 
        # print(words)
        for word in words:  
        #         wordfreq = [words.count(word)]
        #print(dict(list(zip(words,wordfreq))))
            #for word in pat:        
                #print(word)
                #print(len(word))
                if word in d:
                    d[word] = d[word]+1
                else:
                    d[word] = 1 
      
    for key in list(d.keys()): 
    # for wordfreq in d.items():
    #     print(wordfreq) 
            #print(d[key]) 
        #print(d.keys())
        # print(key, ":", d[key])  
        freq = d[key]/count 
        if(freq>0.001):
            tot = tot+1  
            result.update({key:freq})
    print(result)   
    #print(count)
    print(tot)
# Creating a workbook and add a worksheet.
#     workbook = xlsxwriter.Workbook('pride_and_prejudice_word_stats.xlsx')
#     worksheet = workbook.add_worksheet('Word Frequency Stats')
#     row = 0
#     col = 0

# # Iterating over the data and write it out row by row.
#     for key,freq in (result.items()):
#         worksheet.write(row, col,     key)
#         worksheet.write(row, col + 1, freq)
#         row += 1
#     workbook.close()

 
analyze("C:/Users/ionad/OneDrive - University College Cork/Desktop/UCC/SEM 2/CS 6507 - Prog in Python with DS Apps/Assignment/2/pride_and_prejudice.docx")    

